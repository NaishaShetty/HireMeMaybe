"""Resume formatter: converts plain-text resume into PDF or DOCX.

Parses the text into:
  - Candidate name (first non-empty line)
  - Contact line(s) (email / phone / URL patterns near the top)
  - Named sections (EXPERIENCE, EDUCATION, SKILLS, etc.)

Then renders each format with professional styling.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

# ---------------------------------------------------------------------------
# Section detection
# ---------------------------------------------------------------------------

_KNOWN_HEADERS = {
    "summary", "professional summary", "objective", "career objective",
    "profile", "about me",
    "experience", "work experience", "employment history", "work history",
    "professional experience", "career history",
    "education", "academic background", "academic history", "qualifications",
    "skills", "technical skills", "core competencies", "key skills",
    "technologies", "tools", "expertise",
    "projects", "personal projects", "portfolio", "side projects",
    "certifications", "certificates", "licenses", "accreditations",
    "awards", "achievements", "honors", "accomplishments",
    "publications", "research",
    "languages", "spoken languages",
    "interests", "hobbies",
    "volunteering", "volunteer experience", "community involvement",
    "references", "contact", "additional information",
}

_CONTACT_RE = re.compile(
    r"@|linkedin|github|portfolio|mailto|https?://|"
    r"\b\d{3}[\s.\-]\d{3}[\s.\-]\d{4}\b|"
    r"\b[\w.+-]+@[\w-]+\.[a-z]{2,}\b",
    re.IGNORECASE,
)


def _is_section_header(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    # All-caps line of reasonable length (not a sentence)
    if stripped == stripped.upper() and 2 < len(stripped) < 60 and not stripped[-1] == ",":
        return True
    # Known header (with optional trailing colon)
    candidate = stripped.rstrip(":").lower()
    return candidate in _KNOWN_HEADERS


def parse_resume_text(text: str) -> dict[str, Any]:
    """Return {name, contact: [str], sections: OrderedDict[str, [str]]}."""
    lines = text.splitlines()

    name = ""
    contact_lines: list[str] = []
    sections: dict[str, list[str]] = {}
    current_section: str | None = None
    current_content: list[str] = []
    name_found = False
    header_zone_over = False  # once we hit first section, contact zone ends

    for raw_line in lines:
        stripped = raw_line.strip()

        # Grab name from first non-empty line
        if not name_found:
            if stripped:
                name = stripped
                name_found = True
            continue

        if _is_section_header(stripped):
            # Flush previous section
            if current_section is not None:
                sections[current_section] = current_content
            current_section = stripped.rstrip(":").upper()
            current_content = []
            header_zone_over = True
            continue

        if current_section is not None:
            current_content.append(raw_line)
        else:
            # Still in header zone — collect contact info
            if stripped and (not header_zone_over):
                contact_lines.append(stripped)

    # Flush last section
    if current_section is not None:
        sections[current_section] = current_content

    # Deduplicate blank lines inside section content
    cleaned_sections: dict[str, list[str]] = {}
    for sec, content in sections.items():
        cleaned: list[str] = []
        prev_blank = False
        for ln in content:
            is_blank = not ln.strip()
            if is_blank and prev_blank:
                continue
            cleaned.append(ln)
            prev_blank = is_blank
        cleaned_sections[sec] = cleaned

    return {"name": name, "contact": contact_lines, "sections": cleaned_sections}


# ---------------------------------------------------------------------------
# DOCX formatter
# ---------------------------------------------------------------------------

_PURPLE = (0x5B, 0x21, 0xB6)
_DARK   = (0x1A, 0x1A, 0x2E)
_GREY   = (0x44, 0x44, 0x44)


def _add_bottom_border(paragraph: Any, color_hex: str = "5B21B6") -> None:
    """Add a bottom border to a paragraph via raw XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def format_docx(resume_data: dict[str, Any]) -> bytes:
    """Render resume_data as a .docx file and return raw bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.85)
        section.right_margin  = Inches(0.85)

    # Remove default paragraph spacing
    from docx.oxml.ns import qn
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    def _add_para(text: str, bold: bool = False, size: int = 11,
                  color: tuple = _DARK, align=None, space_before: int = 0,
                  space_after: int = 2) -> Any:
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(*color)
        return p

    # ── Name ──────────────────────────────────────────────────────────
    name_para = _add_para(
        resume_data["name"],
        bold=True, size=22, color=_DARK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )

    # ── Contact lines ─────────────────────────────────────────────────
    for line in resume_data["contact"]:
        _add_para(line, size=10, color=_GREY,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    # Horizontal rule after contact block
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(6)
    hr.paragraph_format.space_after  = Pt(4)
    _add_bottom_border(hr, "5B21B6")

    # ── Sections ──────────────────────────────────────────────────────
    for section_name, content_lines in resume_data["sections"].items():
        # Section header
        hdr = _add_para(
            section_name, bold=True, size=12,
            color=_PURPLE, space_before=10, space_after=3,
        )
        _add_bottom_border(hdr, "D1B4F8")

        for raw_line in content_lines:
            stripped = raw_line.strip()
            if not stripped:
                # Small spacer
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(2)
                continue

            is_bullet = stripped[0] in "•-*"
            if is_bullet:
                text = stripped.lstrip("•-* ").strip()
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent  = Inches(0.2)
                p.paragraph_format.space_after  = Pt(2)
                run = p.add_run(text)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(*_GREY)
            else:
                # Detect sub-headers (lines ending with | or containing | separator,
                # or all-bold company/role lines)
                is_subheader = "|" in stripped or stripped.endswith(":")
                p = _add_para(
                    stripped,
                    bold=is_subheader,
                    size=11,
                    color=_DARK if is_subheader else _GREY,
                    space_after=2,
                )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF formatter
# ---------------------------------------------------------------------------

def format_pdf(resume_data: dict[str, Any]) -> bytes:
    """Render resume_data as a PDF file and return raw bytes."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.colors import HexColor, black
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    purple     = HexColor("#5B21B6")
    dark       = HexColor("#1A1A2E")
    grey       = HexColor("#444444")
    light_grey = HexColor("#888888")
    rule_color = HexColor("#D1B4F8")

    name_style = ParagraphStyle(
        "ResumeName",
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=28,
        textColor=dark,
        alignment=TA_CENTER,
        spaceAfter=4,
    )
    contact_style = ParagraphStyle(
        "ResumeContact",
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        textColor=grey,
        alignment=TA_CENTER,
        spaceAfter=2,
    )
    section_style = ParagraphStyle(
        "SectionHeader",
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=16,
        textColor=purple,
        spaceBefore=12,
        spaceAfter=3,
    )
    body_style = ParagraphStyle(
        "Body",
        fontName="Helvetica",
        fontSize=11,
        leading=15,
        textColor=grey,
        spaceAfter=3,
        leftIndent=0,
    )
    subheader_style = ParagraphStyle(
        "SubHeader",
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=15,
        textColor=dark,
        spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        "Bullet",
        fontName="Helvetica",
        fontSize=11,
        leading=15,
        textColor=grey,
        spaceAfter=2,
        leftIndent=16,
        bulletIndent=4,
        bulletFontName="Helvetica",
        bulletFontSize=11,
    )

    story = []

    # Name
    story.append(Paragraph(resume_data["name"], name_style))

    # Contact
    for line in resume_data["contact"]:
        # Escape any XML special chars
        safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe, contact_style))

    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=1.5, color=purple, spaceAfter=4))

    # Sections
    for section_name, content_lines in resume_data["sections"].items():
        story.append(Paragraph(section_name, section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=rule_color, spaceAfter=4))

        for raw_line in content_lines:
            stripped = raw_line.strip()
            if not stripped:
                story.append(Spacer(1, 4))
                continue

            safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

            if stripped[0] in "•-*":
                text = safe.lstrip("•\\-* ").strip()
                story.append(Paragraph(f"• {text}", bullet_style))
            elif "|" in stripped or stripped.endswith(":"):
                story.append(Paragraph(safe, subheader_style))
            else:
                story.append(Paragraph(safe, body_style))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def export_cover_letter_pdf(cover_letter_text: str, company: str = "") -> bytes:
    """Render a plain-text cover letter as a clean PDF and return raw bytes."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        rightMargin=1.0 * inch,
        leftMargin=1.0 * inch,
        topMargin=1.0 * inch,
        bottomMargin=1.0 * inch,
    )

    purple = HexColor("#5B21B6")
    dark   = HexColor("#1A1A2E")
    grey   = HexColor("#444444")

    header_style = ParagraphStyle(
        "CLHeader",
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=15,
        textColor=HexColor("#888888"),
        alignment=TA_CENTER,
        spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "CLBody",
        fontName="Helvetica",
        fontSize=11,
        leading=18,
        textColor=grey,
        spaceAfter=8,
    )
    signature_style = ParagraphStyle(
        "CLSig",
        fontName="Helvetica",
        fontSize=11,
        leading=16,
        textColor=dark,
        spaceAfter=4,
    )

    story = []

    # Thin header bar
    if company:
        safe = company.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(f"Cover Letter — {safe}", header_style))
    else:
        story.append(Paragraph("Cover Letter", header_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=purple, spaceAfter=20))

    # Split into paragraphs
    paragraphs = [p.strip() for p in cover_letter_text.split("\n\n") if p.strip()]
    for i, para in enumerate(paragraphs):
        lines = para.splitlines()
        # Detect signature block (short lines at the end)
        is_sig = i == len(paragraphs) - 1 and len(lines) <= 3
        safe_para = "<br/>".join(
            line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            for line in lines
        )
        style = signature_style if is_sig else body_style
        story.append(Paragraph(safe_para, style))
        if not is_sig:
            story.append(Spacer(1, 4))

    doc.build(story)
    return buf.getvalue()


def export_resume(resume_text: str, fmt: str) -> tuple[bytes, str, str]:
    """Parse *resume_text* and render as *fmt* ('pdf' or 'docx').

    Returns (file_bytes, media_type, filename).
    """
    data = parse_resume_text(resume_text)

    # Derive a safe filename from the candidate's name
    safe_name = re.sub(r"[^a-zA-Z0-9_\- ]", "", data["name"]).strip().replace(" ", "_")
    safe_name = safe_name or "optimized_resume"

    if fmt == "docx":
        content = format_docx(data)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{safe_name}_optimized.docx"
    else:
        content = format_pdf(data)
        media_type = "application/pdf"
        filename = f"{safe_name}_optimized.pdf"

    return content, media_type, filename
